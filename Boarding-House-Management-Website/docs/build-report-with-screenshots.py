"""
build-report-with-screenshots.py
Tạo file Word báo cáo hướng dẫn sử dụng kèm ảnh chụp màn hình
Dùng python-docx để dựng docx trực tiếp với ảnh nhúng vào đúng từng mục
"""

import os
import sys
from pathlib import Path

# ── Install dependencies ───────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────────────────
BASE = Path(__file__).parent
SS = BASE / "screenshots"
OUT = BASE / "docs" / "bao-cao-huong-dan-su-dung-he-thong-with-screenshots.docx"

IMG_W = Inches(5.5)   # Standard image width in doc

# ── Helpers ────────────────────────────────────────────────────────────────

def img(name):
    """Return full path if screenshot exists, else None."""
    p = SS / name
    return str(p) if p.exists() else None


def add_heading(doc, text, level):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_screenshot(doc, filename, caption=""):
    """Add screenshot with optional caption. Skip if file missing."""
    path = img(filename)
    if not path:
        return
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=IMG_W)
        if caption:
            cap = doc.add_paragraph(f"Hình: {caption}")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].italic = True
            cap.runs[0].font.size = Pt(9)
    except Exception as e:
        print(f"  [WARN] Cannot add image {filename}: {e}")


def add_table(doc, headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    # Data rows
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
    doc.add_paragraph()  # spacing after table


def add_page_break(doc):
    doc.add_page_break()


# ── Build document ─────────────────────────────────────────────────────────

doc = Document()

# Default font
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

# ════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run("BÁO CÁO KHẢO SÁT VÀ HƯỚNG DẪN SỬ DỤNG HỆ THỐNG")
run.bold = True
run.font.size = Pt(18)

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_para.add_run("AKDD House — Boarding House Management System")
run2.bold = True
run2.font.size = Pt(14)

doc.add_paragraph()
doc.add_paragraph()

meta = [
    ("Tên hệ thống", "AKDD House — Boarding House Management Website"),
    ("Phiên bản", "1.0"),
    ("Ngày khảo sát", "26/03/2026"),
    ("Phương pháp", "Playwright Automated Browser Exploration"),
    ("URL hệ thống", "http://localhost:9999/House_management1"),
    ("Người thực hiện", "[Tên người thực hiện]"),
    ("Giảng viên hướng dẫn", "[Tên giảng viên]"),
]
add_table(doc, ["Thông tin", "Chi tiết"], meta)

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHƯƠNG 1 — TỔNG QUAN HỆ THỐNG
# ════════════════════════════════════════════════════════════════

add_heading(doc, "1. Tổng Quan Hệ Thống", 1)

add_para(doc, "AKDD House là hệ thống quản lý nhà trọ trực tuyến, số hóa toàn bộ quy trình vận hành: quản lý phòng, hợp đồng, hóa đơn, dịch vụ tiện ích và giao tiếp giữa chủ trọ và khách thuê. Hệ thống phục vụ hai nhóm người dùng: Quản trị viên (Admin/Staff) và Khách thuê (Customer).")

add_heading(doc, "1.1 Thông Tin Kỹ Thuật", 2)
add_table(doc, ["Thành phần", "Thông tin"], [
    ("Backend", "Java EE — Jakarta Servlet 5.0 + JSP"),
    ("Giao diện", "Bootstrap 5.3, Bootstrap Icons"),
    ("Cơ sở dữ liệu", "Microsoft SQL Server 2019+"),
    ("Máy chủ", "Apache Tomcat 10+ / GlassFish"),
    ("Build tool", "Apache Maven"),
])

add_heading(doc, "1.2 Danh Sách Module", 2)
add_table(doc, ["STT", "Nhóm", "Module", "Vai trò"], [
    ("1", "Tài chính", "Hóa đơn (Bills)", "Admin + Customer (xem)"),
    ("2", "Tài chính", "Đặt cọc (Deposits)", "Admin"),
    ("3", "Tài chính", "Bảng giá (Price Categories)", "Admin"),
    ("4", "Tài sản", "Phòng (Rooms)", "Admin + Customer (xem)"),
    ("5", "Tài sản", "Cơ sở vật chất (Facilities)", "Admin"),
    ("6", "Tài sản", "Tiện nghi (Amenities)", "Admin"),
    ("7", "Tài sản", "Điện/Nước/Gas (Utilities)", "Admin"),
    ("8", "Khách thuê", "Hợp đồng (Contracts)", "Admin + Customer (xem)"),
    ("9", "Khách thuê", "Quản lý khách thuê (Customers)", "Admin"),
    ("10", "Dịch vụ", "Danh mục dịch vụ (Services)", "Admin + Customer (đặt)"),
    ("11", "Dịch vụ", "Yêu cầu dịch vụ", "Admin + Customer"),
    ("12", "Hệ thống", "Tài khoản người dùng (Users)", "Admin"),
    ("13", "Hệ thống", "Thông báo (Notifications)", "Admin + Customer (xem)"),
    ("14", "Hệ thống", "Nhật ký hoạt động (Activity Logs)", "Admin — chỉ đọc"),
])

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHƯƠNG 2 — PHƯƠNG PHÁP KHẢO SÁT
# ════════════════════════════════════════════════════════════════

add_heading(doc, "2. Phương Pháp Khảo Sát", 1)
add_para(doc, "Quá trình khảo sát sử dụng Playwright — framework tự động hóa trình duyệt của Microsoft — để duyệt toàn bộ hệ thống theo cả hai vai trò. Toàn bộ thao tác được thực hiện trên trình duyệt Chromium ở chế độ headless.")

add_heading(doc, "2.1 Dữ Liệu Thu Thập", 2)
add_table(doc, ["Module", "Dữ liệu thực tế"], [
    ("Phòng", "10 phòng"),
    ("Hợp đồng", "10 hợp đồng"),
    ("Hóa đơn", "5 hóa đơn"),
    ("Tiền đặt cọc", "13 giao dịch"),
    ("Khách thuê", "10 khách thuê"),
    ("Dịch vụ", "4 loại dịch vụ"),
    ("Yêu cầu dịch vụ", "14 yêu cầu"),
    ("Thông báo", "8 thông báo"),
    ("Bảng giá", "11 danh mục"),
    ("Điện/Nước/Gas", "4 loại tiện ích"),
    ("Tài khoản", "10 người dùng"),
])

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHƯƠNG 3 — ĐĂNG NHẬP VÀ XÁC THỰC
# ════════════════════════════════════════════════════════════════

add_heading(doc, "3. Đăng Nhập Và Xác Thực", 1)

add_heading(doc, "3.1 Trang Đăng Nhập", 2)
add_para(doc, "URL: /auth?action=login — Tiêu đề trang: \"Login - AKDD House\"")
add_screenshot(doc, "01-admin-login.png", "Trang đăng nhập hệ thống")
add_table(doc, ["Trường", "Loại", "Bắt buộc", "Mô tả"], [
    ("Username", "Text", "Có", "Tên đăng nhập tài khoản"),
    ("Password", "Password", "Có", "Mật khẩu (ẩn ký tự)"),
])
add_para(doc, "Sau đăng nhập: Admin/Staff → Admin Dashboard | Customer → Customer Dashboard | Sai thông tin → hiển thị lỗi, ở lại trang đăng nhập.")

add_heading(doc, "3.2 Đăng Ký Tài Khoản", 2)
add_para(doc, "URL: /auth?action=register")
add_screenshot(doc, "05-guest-dang-ky.png", "Trang đăng ký tài khoản khách thuê")

add_heading(doc, "3.3 Khôi Phục Mật Khẩu (OTP Email)", 2)
add_para(doc, "URL: /auth?action=forgetPassword — Quy trình 4 bước: Nhập email → Nhận OTP → Xác minh OTP → Đặt mật khẩu mới.")

add_heading(doc, "3.4 Đổi Mật Khẩu", 2)
add_screenshot(doc, "41-user-change-password.png", "Trang đổi mật khẩu")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHƯƠNG 4 — PHÂN QUYỀN
# ════════════════════════════════════════════════════════════════

add_heading(doc, "4. Phân Quyền Người Dùng", 1)
add_table(doc, ["Vai trò", "Mô tả", "Giao diện sau đăng nhập"], [
    ("admin", "Quản trị viên toàn quyền", "Admin Dashboard"),
    ("staff", "Nhân viên quản lý", "Admin Dashboard"),
    ("customer", "Khách thuê phòng", "Customer Dashboard"),
])

add_heading(doc, "4.1 Kiểm Tra Truy Cập Thực Tế", 2)
add_para(doc, "Kết quả khi tài khoản customer cố tình truy cập trang admin:")
add_table(doc, ["URL truy cập", "Kết quả"], [
    ("/admin?action=dashboard", "Redirect → /guest/dashboard.jsp"),
    ("/bill?action=list", "Redirect → /dashboard (Customer)"),
    ("/manage-customer?action=list", "Redirect → /dashboard"),
    ("/user?action=list", "Redirect → /dashboard"),
    ("/activity-log", "Redirect → /dashboard"),
    ("/utility?action=list", "Redirect → /dashboard"),
    ("/contract?action=list", "Redirect → /contract?action=mycontract"),
])

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHƯƠNG 5 — DASHBOARD
# ════════════════════════════════════════════════════════════════

add_heading(doc, "5. Trang Tổng Quan (Dashboard)", 1)

add_heading(doc, "5.1 Admin Dashboard", 2)
add_screenshot(doc, "02-admin-dashboard.png", "Trang tổng quan Admin")
add_table(doc, ["Chỉ số", "Mô tả"], [
    ("Total Rooms", "Tổng số phòng trong hệ thống"),
    ("Active Contracts", "Số hợp đồng đang còn hiệu lực"),
    ("Unpaid Bills", "Số hóa đơn chưa thanh toán"),
    ("Customers", "Tổng số khách thuê hoạt động"),
])

add_heading(doc, "5.2 Customer Dashboard", 2)
add_screenshot(doc, "31-user-dashboard.png", "Trang tổng quan Customer (tài khoản: sale)")
add_para(doc, "Tài khoản 'sale' đang thuê phòng 106 (Double Room), 1 người ở cùng. Dashboard hiển thị: thông tin phòng, hóa đơn hiện tại, thông báo mới nhất, yêu cầu dịch vụ gần đây.")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHƯƠNG 6 — QUẢN LÝ PHÒNG
# ════════════════════════════════════════════════════════════════

add_heading(doc, "6. Module Quản Lý Phòng (Rooms)", 1)
add_para(doc, "Phạm vi: Admin / Staff | URL: /room?action=list")

add_heading(doc, "6.1 Danh Sách Phòng", 2)
add_screenshot(doc, "crud-room-01-list.png", "Danh sách phòng — Admin")
add_table(doc, ["Cột", "Mô tả"], [
    ("Room Number", "Mã số phòng"),
    ("Status", "Available (xanh) / Occupied (đỏ) / Maintenance (vàng)"),
    ("Action", "Detail / Edit / Delete"),
])

add_heading(doc, "6.2 Xem Chi Tiết Phòng", 2)
add_screenshot(doc, "18-admin-room-detail.png", "Chi tiết phòng — Phòng 101 (Occupied)")

add_heading(doc, "6.3 Thêm Phòng Mới", 2)
add_screenshot(doc, "crud-room-02-create-form.png", "Form thêm phòng mới")
add_table(doc, ["Trường", "Bắt buộc", "Mô tả"], [
    ("Room Number", "Có", "Mã phòng (101, A2-03...)"),
    ("Category", "Có", "Danh mục phòng (liên kết bảng giá)"),
    ("Status", "Không", "Trạng thái khởi tạo (mặc định: available)"),
    ("Image", "Không", "Đường dẫn hình ảnh"),
])

add_heading(doc, "6.4 Chỉnh Sửa Phòng", 2)
add_screenshot(doc, "crud-room-03-edit-form.png", "Form chỉnh sửa phòng")

add_heading(doc, "6.5 Customer — Duyệt Phòng", 2)
add_screenshot(doc, "40-user-room-categories.png", "Trang duyệt phòng theo danh mục (Customer)")
add_screenshot(doc, "43-user-room-detail.png", "Chi tiết phòng và Booking Request (Customer)")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHƯƠNG 7 — CƠ SỞ VẬT CHẤT
# ════════════════════════════════════════════════════════════════

add_heading(doc, "7. Module Quản Lý Cơ Sở Vật Chất (Facilities)", 1)
add_para(doc, "Phạm vi: Admin / Staff | URL: /facility?action=list")

add_heading(doc, "7.1 Danh Sách Cơ Sở Vật Chất", 2)
add_screenshot(doc, "07-admin-facilities-list.png", "Danh sách cơ sở vật chất")

add_heading(doc, "7.2 Thêm Mới", 2)
add_screenshot(doc, "07-admin-facilities-create.png", "Form thêm cơ sở vật chất")
add_table(doc, ["Trường", "Bắt buộc", "Mô tả"], [
    ("Facility Name", "Có", "Tên hạng mục"),
    ("Category", "Có", "Mã danh mục"),
    ("Description", "Không", "Mô tả chi tiết"),
    ("Image", "Không", "Đường dẫn hình ảnh"),
])

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHƯƠNG 8 — TIỆN NGHI
# ════════════════════════════════════════════════════════════════

add_heading(doc, "8. Module Quản Lý Tiện Nghi (Amenities)", 1)
add_para(doc, "Phạm vi: Admin / Staff | URL: /amenity?action=list | Hỗ trợ ẩn/hiện (soft delete)")

add_heading(doc, "8.1 Danh Sách Tiện Nghi", 2)
add_screenshot(doc, "08-admin-amenities-list.png", "Danh sách tiện nghi")

add_heading(doc, "8.2 Thêm Tiện Nghi Mới", 2)
add_screenshot(doc, "08-admin-amenities-create.png", "Form thêm tiện nghi")
add_table(doc, ["Trường", "Bắt buộc"], [
    ("Amenity Name", "Có"),
    ("Description", "Không"),
])

add_para(doc, "Thao tác Hide: ẩn tiện nghi khỏi danh sách chọn khi gán phòng. Thao tác Restore: khôi phục tiện nghi đã ẩn.")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHƯƠNG 9 — ĐIỆN/NƯỚC/GAS
# ════════════════════════════════════════════════════════════════

add_heading(doc, "9. Module Quản Lý Điện — Nước — Gas (Utilities)", 1)
add_para(doc, "Phạm vi: Admin / Staff | URL: /utility?action=list | 4 loại tiện ích đang hoạt động.")

add_heading(doc, "9.1 Danh Sách Tiện Ích", 2)
add_screenshot(doc, "crud-utility-01-list.png", "Danh sách loại điện/nước/gas")

add_heading(doc, "9.2 Thêm Loại Tiện Ích", 2)
add_table(doc, ["Trường", "Bắt buộc", "Mô tả"], [
    ("Utility Name", "Có", "Electricity, Water, Gas..."),
    ("Unit", "Có", "kWh, m³, kg..."),
    ("Description", "Không", "Mô tả thêm"),
])

add_heading(doc, "9.3 Quản Lý Bảng Giá Tiện Ích", 2)
add_screenshot(doc, "crud-utility-04-add-price-form.png", "Form thêm mức giá tiện ích")

add_heading(doc, "9.4 Ghi Nhận Chỉ Số Tiêu Thụ", 2)
add_screenshot(doc, "crud-utility-03-add-usage-form.png", "Form ghi chỉ số tiêu thụ")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHƯƠNG 10 — BẢNG GIÁ
# ════════════════════════════════════════════════════════════════

add_heading(doc, "10. Module Bảng Giá Phòng (Price Categories)", 1)
add_para(doc, "Phạm vi: Admin / Staff | URL: /price?action=categories | 11 danh mục giá.")

add_heading(doc, "10.1 Danh Sách Danh Mục Giá", 2)
add_screenshot(doc, "crud-roomcat-01-list.png", "Danh sách bảng giá phòng")

add_heading(doc, "10.2 Thêm Danh Mục Giá", 2)
add_screenshot(doc, "crud-roomcat-02-create-form.png", "Form thêm danh mục giá")
add_table(doc, ["Trường", "Bắt buộc", "Mô tả"], [
    ("Category Code", "Có", "SINGLE, DOUBLE, STUDIO..."),
    ("Category Type", "Có", "Loại phòng"),
    ("Unit", "Không", "Đơn vị tính giá"),
])

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHƯƠNG 11 — HỢP ĐỒNG
# ════════════════════════════════════════════════════════════════

add_heading(doc, "11. Module Quản Lý Hợp Đồng (Contracts)", 1)
add_para(doc, "Phạm vi: Admin/Staff (quản lý toàn bộ) | Customer (xem hợp đồng của mình) | URL: /contract?action=list")

add_heading(doc, "11.1 Danh Sách Hợp Đồng", 2)
add_screenshot(doc, "crud-contract-01-list.png", "Danh sách hợp đồng — 10 hợp đồng")
add_table(doc, ["Cột", "Mô tả"], [
    ("ID", "Mã hợp đồng"),
    ("Room", "Phòng được thuê"),
    ("Primary Tenant", "Khách thuê chính"),
    ("Start Date / End Date", "Thời hạn hợp đồng"),
    ("Deposit", "Tiền đặt cọc"),
    ("Tenants", "Số người trong hợp đồng"),
    ("Status", "Active / Terminated / Expired"),
])

add_heading(doc, "11.2 Tạo Hợp Đồng Mới", 2)
add_screenshot(doc, "crud-contract-02-create-form.png", "Form tạo hợp đồng — Bước 1: Chọn phòng")
add_table(doc, ["Trường", "Bắt buộc", "Mô tả"], [
    ("Phòng", "Có", "Chọn từ grid phòng Available"),
    ("Primary Tenant", "Có", "Khách thuê chính"),
    ("Start Date", "Có", "Ngày bắt đầu"),
    ("End Date", "Không", "Bỏ trống = không thời hạn"),
    ("Deposit", "Có", "Tiền đặt cọc (VND)"),
    ("Tenant Name", "Có", "Họ tên người thuê thêm"),
    ("CCCD/CMND", "Không", "Số chứng minh nhân dân"),
])

add_heading(doc, "11.3 Chi Tiết Hợp Đồng", 2)
add_screenshot(doc, "crud-contract-03-detail.png", "Chi tiết hợp đồng")

add_heading(doc, "11.4 Chỉnh Sửa Hợp Đồng", 2)
add_screenshot(doc, "crud-contract-05-edit-form.png", "Form chỉnh sửa hợp đồng")

add_heading(doc, "11.5 Customer — Hợp Đồng Của Tôi", 2)
add_screenshot(doc, "34-user-contracts.png", "Trang hợp đồng của khách thuê (My Contracts)")
add_screenshot(doc, "35-user-sign-contract.png", "Trang ký hợp đồng mới (Sign Contract)")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHƯƠNG 12 — KHÁCH THUÊ
# ════════════════════════════════════════════════════════════════

add_heading(doc, "12. Module Quản Lý Khách Thuê (Customers)", 1)
add_para(doc, "Phạm vi: Admin / Staff | URL: /manage-customer?action=list | 10 khách thuê.")

add_heading(doc, "12.1 Danh Sách Khách Thuê", 2)
add_screenshot(doc, "11-admin-customers-list.png", "Danh sách khách thuê")

add_heading(doc, "12.2 Thêm Khách Thuê Mới", 2)
add_screenshot(doc, "11-admin-customers-create.png", "Form thêm khách thuê")
add_table(doc, ["Trường", "Bắt buộc"], [
    ("Username", "Có"),
    ("Password", "Có"),
    ("Full Name", "Có"),
    ("Email", "Không"),
    ("Phone", "Không"),
])

add_heading(doc, "12.3 Chi Tiết Khách Thuê", 2)
add_screenshot(doc, "20-admin-customer-detail.png", "Chi tiết hồ sơ khách thuê")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHƯƠNG 13 — HÓA ĐƠN
# ════════════════════════════════════════════════════════════════

add_heading(doc, "13. Module Quản Lý Hóa Đơn (Bills)", 1)
add_para(doc, "Phạm vi: Admin/Staff (tạo, quản lý) | Customer (xem hóa đơn của mình) | URL: /bill?action=list | 5 hóa đơn.")

add_heading(doc, "13.1 Danh Sách Hóa Đơn — Admin", 2)
add_screenshot(doc, "crud-bill-01-list.png", "Danh sách hóa đơn — Admin")
add_table(doc, ["Trạng thái", "Mô tả"], [
    ("Unpaid", "Chưa thanh toán"),
    ("Paid", "Đã thanh toán"),
    ("Overdue", "Quá hạn thanh toán"),
])

add_heading(doc, "13.2 Tạo Hóa Đơn Mới", 2)
add_screenshot(doc, "crud-bill-02-create-form.png", "Form tạo hóa đơn mới (hỗ trợ nhiều dòng mục)")
add_table(doc, ["Trường", "Bắt buộc", "Mô tả"], [
    ("Contract", "Có", "Chọn hợp đồng"),
    ("Period", "Có", "Kỳ thanh toán"),
    ("Due Date", "Có", "Hạn thanh toán"),
    ("Description (item)", "Có", "Room Rent, Electricity, Water..."),
    ("Quantity", "Có", "Số lượng (mặc định: 1)"),
    ("Unit Price", "Có", "Đơn giá (VND)"),
])

add_heading(doc, "13.3 Chi Tiết Hóa Đơn", 2)
add_screenshot(doc, "crud-bill-03-detail.png", "Chi tiết hóa đơn với các khoản mục")

add_heading(doc, "13.4 Cập Nhật Trạng Thái Thanh Toán", 2)
add_screenshot(doc, "crud-bill-04-status-update.png", "Cập nhật trạng thái hóa đơn")

add_heading(doc, "13.5 Customer — Hóa Đơn Của Tôi", 2)
add_screenshot(doc, "33-user-bills.png", "Trang hóa đơn của khách thuê (My Bills)")
add_screenshot(doc, "crud-cust-04-bill-detail.png", "Chi tiết hóa đơn — phía khách thuê")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHƯƠNG 14 — TIỀN ĐẶT CỌC
# ════════════════════════════════════════════════════════════════

add_heading(doc, "14. Module Quản Lý Tiền Đặt Cọc (Deposits)", 1)
add_para(doc, "Phạm vi: Admin / Staff | URL: /deposit?action=all | 13 giao dịch.")

add_heading(doc, "14.1 Danh Sách Giao Dịch Đặt Cọc", 2)
add_screenshot(doc, "crud-deposit-01-list.png", "Danh sách giao dịch đặt cọc")

add_heading(doc, "14.2 Thêm Giao Dịch Mới", 2)
add_screenshot(doc, "crud-deposit-02-create-form.png", "Form giao dịch đặt cọc")
add_table(doc, ["Loại giao dịch", "Mô tả"], [
    ("Deposit", "Nạp tiền cọc khi ký hợp đồng"),
    ("Refund", "Hoàn trả khi kết thúc hợp đồng"),
    ("Deduction", "Khấu trừ (hư hỏng, vi phạm...)"),
])

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHƯƠNG 15 — DỊCH VỤ
# ════════════════════════════════════════════════════════════════

add_heading(doc, "15. Module Quản Lý Dịch Vụ (Services)", 1)
add_para(doc, "Phạm vi: Admin/Staff (quản lý) | Customer (đặt và theo dõi) | 4 dịch vụ | 14 yêu cầu đang xử lý.")

add_heading(doc, "15.1 Danh Sách Dịch Vụ — Admin", 2)
add_screenshot(doc, "crud-service-01-admin-list.png", "Danh sách dịch vụ — Admin")

add_heading(doc, "15.2 Thêm Dịch Vụ Mới", 2)
add_screenshot(doc, "crud-service-02-create-form.png", "Form thêm dịch vụ")

add_heading(doc, "15.3 Quản Lý Yêu Cầu Dịch Vụ", 2)
add_screenshot(doc, "crud-service-03-request-list.png", "Danh sách yêu cầu dịch vụ từ khách thuê")
add_table(doc, ["Thao tác", "Mô tả"], [
    ("Approve", "Chấp thuận yêu cầu"),
    ("Reject", "Từ chối kèm lý do"),
    ("Mark Billed", "Đánh dấu đã tính phí vào hóa đơn"),
])

add_heading(doc, "15.4 Customer — Đặt Dịch Vụ", 2)
add_screenshot(doc, "38-user-request-service.png", "Form đặt dịch vụ — Customer")

add_heading(doc, "15.5 Customer — Lịch Sử Dịch Vụ", 2)
add_screenshot(doc, "37-user-service-history.png", "Lịch sử yêu cầu dịch vụ — Customer")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHƯƠNG 16 — THÔNG BÁO
# ════════════════════════════════════════════════════════════════

add_heading(doc, "16. Module Quản Lý Thông Báo (Notifications)", 1)
add_para(doc, "Phạm vi: Admin/Staff (tạo, quản lý) | Customer (xem) | 8 thông báo.")

add_heading(doc, "16.1 Danh Sách Thông Báo — Admin", 2)
add_screenshot(doc, "crud-notif-01-list.png", "Danh sách thông báo — Admin")

add_heading(doc, "16.2 Tạo Thông Báo Mới", 2)
add_screenshot(doc, "crud-notif-02-create-form.png", "Form tạo thông báo mới")
add_table(doc, ["Trường", "Bắt buộc", "Mô tả"], [
    ("Title", "Có", "Tiêu đề thông báo"),
    ("Content", "Có", "Nội dung đầy đủ"),
    ("Target Contract", "Không", "Để trống = gửi tất cả; điền ID = gửi riêng"),
])
add_screenshot(doc, "crud-notif-03-broadcast-filled.png", "Thông báo gửi toàn bộ khách thuê")
add_screenshot(doc, "crud-notif-05-targeted-filled.png", "Thông báo gửi riêng theo hợp đồng")

add_heading(doc, "16.3 Chi Tiết & Chỉnh Sửa Thông Báo", 2)
add_screenshot(doc, "crud-notif-06-detail.png", "Chi tiết thông báo")
add_screenshot(doc, "crud-notif-07-edit-form.png", "Form chỉnh sửa thông báo")

add_heading(doc, "16.4 Customer — Xem Thông Báo", 2)
add_screenshot(doc, "39-user-notifications.png", "Trang thông báo — Customer")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHƯƠNG 17 — QUẢN LÝ TÀI KHOẢN
# ════════════════════════════════════════════════════════════════

add_heading(doc, "17. Module Quản Lý Tài Khoản Người Dùng (Users)", 1)
add_para(doc, "Phạm vi: Admin / Staff | URL: /user?action=list | Giao diện tiếng Việt | 10 tài khoản.")

add_heading(doc, "17.1 Danh Sách Người Dùng", 2)
add_screenshot(doc, "14-admin-users-list.png", "Danh sách tài khoản người dùng")
add_table(doc, ["Cột", "Mô tả"], [
    ("Người dùng", "Tên đăng nhập và họ tên"),
    ("Email / Điện thoại", "Thông tin liên lạc"),
    ("Vai trò", "admin / staff / customer"),
    ("Thao tác", "Chỉnh sửa / Xóa"),
])

add_heading(doc, "17.2 Chỉnh Sửa Tài Khoản", 2)
add_screenshot(doc, "17-admin-edit-user.png", "Form chỉnh sửa tài khoản")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHƯƠNG 18 — NHẬT KÝ HOẠT ĐỘNG
# ════════════════════════════════════════════════════════════════

add_heading(doc, "18. Module Nhật Ký Hoạt Động (Activity Logs)", 1)
add_para(doc, "Phạm vi: Admin / Staff — CHỈ ĐỌC | URL: /activity-log")
add_para(doc, "Ghi lại toàn bộ hành động của người dùng để kiểm tra, kiểm toán nội bộ.")

add_screenshot(doc, "16-admin-activity-log.png", "Trang nhật ký hoạt động")
add_table(doc, ["Bộ lọc", "Mô tả"], [
    ("Search", "Tìm theo từ khóa"),
    ("Type", "Loại hành động"),
    ("Date From / To", "Khoảng thời gian"),
    ("User", "Lọc theo tài khoản cụ thể"),
])

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHƯƠNG 19 — CHỨC NĂNG CUSTOMER
# ════════════════════════════════════════════════════════════════

add_heading(doc, "19. Chức Năng Dành Riêng Cho Khách Thuê", 1)

add_heading(doc, "19.1 Hồ Sơ Cá Nhân", 2)
add_screenshot(doc, "32-user-profile.png", "Trang hồ sơ cá nhân")
add_screenshot(doc, "46-user-edit-profile.png", "Form chỉnh sửa hồ sơ (fullName, email, phone)")

add_heading(doc, "19.2 Hóa Đơn Của Tôi", 2)
add_screenshot(doc, "33-user-bills.png", "Danh sách hóa đơn — Customer")

add_heading(doc, "19.3 Hợp Đồng Của Tôi", 2)
add_screenshot(doc, "34-user-contracts.png", "Thông tin hợp đồng — Customer")

add_heading(doc, "19.4 Dịch Vụ", 2)
add_screenshot(doc, "36-user-services.png", "Danh sách dịch vụ có sẵn")
add_screenshot(doc, "38-user-request-service.png", "Form đặt dịch vụ")
add_screenshot(doc, "37-user-service-history.png", "Lịch sử dịch vụ")

add_heading(doc, "19.5 Duyệt Phòng", 2)
add_screenshot(doc, "40-user-room-categories.png", "Duyệt phòng theo danh mục")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHƯƠNG 20 — TỔNG HỢP PHÂN QUYỀN
# ════════════════════════════════════════════════════════════════

add_heading(doc, "20. Tổng Hợp Phân Quyền", 1)

add_heading(doc, "20.1 Bảng So Sánh Admin — Customer", 2)
add_table(doc, ["Chức năng", "Admin", "Customer"], [
    ("Admin Dashboard", "Có", "Không (redirect)"),
    ("Customer Dashboard", "Không (redirect)", "Có"),
    ("Quản lý phòng (CRUD)", "Có", "Chỉ xem"),
    ("Quản lý tiện nghi", "Có", "Không"),
    ("Quản lý điện/nước/gas", "Có", "Không"),
    ("Bảng giá phòng", "Có", "Không"),
    ("Tạo & quản lý hợp đồng", "Có", "Gửi yêu cầu ký"),
    ("Xem tất cả hợp đồng", "Có", "Chỉ của mình"),
    ("Quản lý khách thuê", "Có", "Không"),
    ("Tạo hóa đơn", "Có", "Không"),
    ("Xem tất cả hóa đơn", "Có", "Chỉ của mình"),
    ("Quản lý đặt cọc", "Có", "Không"),
    ("Quản lý dịch vụ (CRUD)", "Có", "Không"),
    ("Đặt và theo dõi dịch vụ", "Không trực tiếp", "Có"),
    ("Duyệt yêu cầu dịch vụ", "Có", "Không"),
    ("Tạo thông báo", "Có", "Không"),
    ("Xem thông báo", "Có", "Chỉ của mình"),
    ("Quản lý tài khoản (Users)", "Có", "Không"),
    ("Nhật ký hoạt động", "Có (chỉ đọc)", "Không"),
    ("Đổi mật khẩu", "Có", "Có"),
    ("Cập nhật hồ sơ", "Qua module Users", "Qua trang Profile"),
])

add_heading(doc, "20.2 URL Tham Chiếu Nhanh — Admin", 2)
add_table(doc, ["Chức năng", "URL"], [
    ("Dashboard", "/dashboard"),
    ("Danh sách phòng", "/room?action=list"),
    ("Tạo phòng mới", "/room?action=create"),
    ("Danh sách hợp đồng", "/contract?action=list"),
    ("Tạo hợp đồng", "/contract?action=create"),
    ("Danh sách hóa đơn", "/bill?action=list"),
    ("Tạo hóa đơn", "/bill?action=create"),
    ("Tiền đặt cọc", "/deposit?action=all"),
    ("Bảng giá", "/price?action=categories"),
    ("Điện/nước/gas", "/utility?action=list"),
    ("Tiện nghi", "/amenity?action=list"),
    ("Cơ sở vật chất", "/facility?action=list"),
    ("Danh sách dịch vụ", "/services?action=adminList"),
    ("Yêu cầu dịch vụ", "/services?action=manageRequests"),
    ("Khách thuê", "/manage-customer?action=list"),
    ("Thông báo", "/notification?action=list"),
    ("Tài khoản", "/user?action=list"),
    ("Nhật ký", "/activity-log"),
])

add_heading(doc, "20.3 URL Tham Chiếu Nhanh — Customer", 2)
add_table(doc, ["Chức năng", "URL"], [
    ("Dashboard", "/dashboard"),
    ("Hồ sơ cá nhân", "/customer?action=profile"),
    ("Chỉnh sửa hồ sơ", "/customer?action=editProfile"),
    ("Hóa đơn của tôi", "/bill?action=mybill"),
    ("Hợp đồng của tôi", "/contract?action=mycontract"),
    ("Ký hợp đồng", "/contract?action=signContract"),
    ("Duyệt phòng", "/room?action=categories"),
    ("Đặt dịch vụ", "/services?action=requestForm"),
    ("Lịch sử dịch vụ", "/services?action=myHistory"),
    ("Thông báo", "/notification?action=publicList"),
    ("Đổi mật khẩu", "/auth?action=changePassword"),
    ("Đăng xuất", "/auth?action=logout"),
])

# ── Save ───────────────────────────────────────────────────────────────────
doc.save(str(OUT))
print(f"[OK] Saved: {OUT}")
print(f"[OK] Size: {OUT.stat().st_size / 1024:.0f} KB")
