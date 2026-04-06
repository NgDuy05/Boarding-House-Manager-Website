"""
build-docx-with-screenshots.py
Builds a professional Word document (user-manual-with-screenshots.docx)
by combining structured markdown content with actual Playwright screenshots.
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS_DIR   = Path(__file__).parent
SHOTS_DIR  = DOCS_DIR / "screenshots"
OUTPUT     = DOCS_DIR / "user-manual-with-screenshots.docx"

# ── Colour palette ───────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1a, 0x3a, 0x5c)
BLUE   = RGBColor(0x2c, 0x5f, 0x8a)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY  = RGBColor(0xF5, 0xF8, 0xFC)
DGRAY  = RGBColor(0x44, 0x44, 0x44)
BLACK  = RGBColor(0x11, 0x11, 0x11)

# ── Helper: shade a table cell ───────────────────────────────────────────────
def shade_cell(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ── Helper: insert a screenshot ──────────────────────────────────────────────
def add_screenshot(doc, filename: str, caption: str, width=Inches(5.5)):
    path = SHOTS_DIR / filename
    if not path.exists():
        # Try partial match
        matches = sorted(SHOTS_DIR.glob(f"*{filename.split('-', 1)[-1] if '-' in filename else filename}*"))
        if matches:
            path = matches[0]
        else:
            p = doc.add_paragraph(f"[Screenshot not found: {filename}]")
            p.runs[0].font.color.rgb = RGBColor(0xAA, 0, 0)
            return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=width)

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size   = Pt(9)
    cap.runs[0].font.italic = True
    cap.runs[0].font.color.rgb = DGRAY
    doc.add_paragraph()   # spacing

# ── Helper: add a styled heading ─────────────────────────────────────────────
def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = NAVY
    p.runs[0].font.size = Pt(18)
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = NAVY
    p.runs[0].font.size = Pt(14)
    return p

def h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = BLUE
    p.runs[0].font.size = Pt(12)
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    return p

def step(doc, text):
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.space_after = Pt(3)
    return p

# ── Helper: simple 2-col info table ─────────────────────────────────────────
def info_table(doc, rows: list[tuple], col_widths=(Inches(2.2), Inches(4.0))):
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, txt in enumerate(['Thông Tin', 'Giá Trị']):
        hdr[i].text = txt
        hdr[i].paragraphs[0].runs[0].font.bold  = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        shade_cell(hdr[i], '1A3A5C')
    for key, val in rows:
        r = table.add_row().cells
        r[0].text = key
        r[1].text = val
        r[0].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()

# ── Build document ───────────────────────────────────────────────────────────
def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── COVER PAGE ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TÀI LIỆU HƯỚNG DẪN SỬ DỤNG\nHỆ THỐNG QUẢN LÝ NHÀ TRỌ")
    run.font.size  = Pt(24)
    run.font.bold  = True
    run.font.color.rgb = NAVY

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Boarding House Management Website")
    r2.font.size  = Pt(14)
    r2.font.color.rgb = BLUE

    doc.add_paragraph()
    doc.add_paragraph()
    info_table(doc, [
        ("Phiên bản",         "1.0"),
        ("Ngày lập",          "25/03/2026"),
        ("Nền tảng",          "Jakarta EE 9.1 · Java 11 · MS SQL Server"),
        ("Giảng viên",        "[Tên giảng viên]"),
        ("Nhóm thực hiện",    "[Tên nhóm]"),
    ])
    doc.add_page_break()

    # ── 1. GIỚI THIỆU HỆ THỐNG ──────────────────────────────────────────────
    h1(doc, "1. Giới Thiệu Hệ Thống")
    body(doc,
        "Hệ thống Quản Lý Nhà Trọ là ứng dụng web số hóa toàn bộ nghiệp vụ quản lý "
        "nhà trọ: phòng trọ, hợp đồng thuê, hóa đơn, tiện ích (điện/nước/gas), dịch vụ "
        "bảo trì và thông báo. Hệ thống phục vụ ba nhóm người dùng: Khách vãng lai, "
        "Người thuê (Customer) và Quản trị viên (Admin)."
    )

    h2(doc, "1.1 Trang Chủ (Khách Vãng Lai)")
    add_screenshot(doc, "01-guest-trang-chu.png", "Hình 1.1 – Trang chủ hiển thị thông tin nhà trọ")

    h2(doc, "1.2 Danh Sách Phòng")
    add_screenshot(doc, "02-guest-danh-sach-phong.png", "Hình 1.2 – Danh sách phòng công khai")
    doc.add_page_break()

    # ── 2. ĐĂNG KÝ / ĐĂNG NHẬP ─────────────────────────────────────────────
    h1(doc, "2. Đăng Ký và Đăng Nhập")

    h2(doc, "2.1 Trang Đăng Nhập")
    body(doc, "Truy cập /auth để đăng nhập. Nhập tên đăng nhập và mật khẩu, nhấn Đăng nhập. "
              "Hệ thống tự chuyển hướng theo vai trò: Admin → /admin/dashboard, Customer → /customer/dashboard.")
    add_screenshot(doc, "04-guest-dang-nhap.png", "Hình 2.1 – Màn hình đăng nhập")

    h2(doc, "2.2 Nhập Thông Tin Đăng Nhập")
    add_screenshot(doc, "06-admin-login-nhap-thong-tin.png", "Hình 2.2 – Nhập tài khoản và mật khẩu")

    h2(doc, "2.3 Trang Đăng Ký")
    body(doc, "Nhấn 'Đăng ký' tại trang đăng nhập để tạo tài khoản mới. "
              "Điền đầy đủ: tên đăng nhập, mật khẩu, họ tên, email, số điện thoại.")
    add_screenshot(doc, "05-guest-dang-ky.png", "Hình 2.3 – Màn hình đăng ký tài khoản")
    doc.add_page_break()

    # ── 3. QUẢN TRỊ VIÊN ───────────────────────────────────────────────────
    h1(doc, "3. Hướng Dẫn Quản Trị Viên (Admin)")
    body(doc, "Sau khi đăng nhập bằng tài khoản admin, hệ thống chuyển đến trang Dashboard tổng quan.")

    h2(doc, "3.1 Dashboard Tổng Quan")
    add_screenshot(doc, "07-admin-dashboard.png", "Hình 3.1 – Dashboard admin: thống kê tổng quan hệ thống")

    h2(doc, "3.2 Quản Lý Phòng")
    h3(doc, "3.2.1 Danh Sách Phòng")
    body(doc, "Hiển thị toàn bộ phòng với trạng thái: Còn trống / Đang sử dụng / Bảo trì. "
              "Cho phép tìm kiếm, lọc theo danh mục và thao tác CRUD.")
    add_screenshot(doc, "08-admin-quan-ly-phong.png", "Hình 3.2 – Danh sách quản lý phòng")

    h3(doc, "3.2.2 Thêm Phòng Mới")
    body(doc, "Nhấn 'Thêm phòng', điền số phòng, chọn danh mục, trạng thái và tải ảnh.")
    add_screenshot(doc, "09-admin-them-phong.png", "Hình 3.3 – Form thêm phòng mới")

    h3(doc, "3.2.3 Danh Mục Phòng")
    body(doc, "Quản lý loại phòng (Standard, Deluxe, Suite...) và giá cơ bản tương ứng.")
    add_screenshot(doc, "10-admin-danh-muc-phong.png", "Hình 3.4 – Quản lý danh mục phòng")
    doc.add_page_break()

    h2(doc, "3.3 Quản Lý Hợp Đồng")
    h3(doc, "3.3.1 Danh Sách Hợp Đồng")
    add_screenshot(doc, "11-admin-hop-dong.png", "Hình 3.5 – Danh sách hợp đồng thuê phòng")

    h3(doc, "3.3.2 Chi Tiết Hợp Đồng")
    body(doc, "Xem đầy đủ thông tin hợp đồng: phòng, người thuê, ngày bắt đầu/kết thúc, "
              "danh sách người cùng thuê (co-tenants) và lịch sử.")
    add_screenshot(doc, "11b-admin-chi-tiet-hop-dong.png", "Hình 3.6 – Chi tiết hợp đồng và người thuê")

    h3(doc, "3.3.3 Tạo Hợp Đồng Mới")
    body(doc, "Chọn phòng, người dùng, nhập ngày bắt đầu/kết thúc và tổng giá thuê.")
    add_screenshot(doc, "12-admin-tao-hop-dong.png", "Hình 3.7 – Form tạo hợp đồng mới")
    doc.add_page_break()

    h2(doc, "3.4 Quản Lý Hóa Đơn")
    h3(doc, "3.4.1 Danh Sách Hóa Đơn")
    add_screenshot(doc, "13-admin-hoa-don.png", "Hình 3.8 – Danh sách hóa đơn")

    h3(doc, "3.4.2 Chi Tiết Hóa Đơn")
    body(doc, "Xem từng khoản mục: tiền phòng, điện, nước, dịch vụ kèm số lượng và đơn giá.")
    add_screenshot(doc, "13b-admin-chi-tiet-hoa-don.png", "Hình 3.9 – Chi tiết các khoản trong hóa đơn")

    h3(doc, "3.4.3 Tạo Hóa Đơn Mới")
    body(doc, "Chọn hợp đồng, thêm khoản mục, đặt ngày đáo hạn và lưu.")
    add_screenshot(doc, "14-admin-tao-hoa-don.png", "Hình 3.10 – Form tạo hóa đơn mới")
    doc.add_page_break()

    h2(doc, "3.5 Quản Lý Tiện Ích")
    body(doc, "Ghi chỉ số đồng hồ điện/nước/gas, quản lý bảng giá theo thời kỳ và theo dõi lượng tiêu thụ.")
    add_screenshot(doc, "15-admin-tien-ich.png", "Hình 3.11 – Quản lý tiện ích điện, nước, gas")

    h2(doc, "3.6 Quản Lý Dịch Vụ")
    body(doc, "Danh mục dịch vụ bảo trì/sửa chữa và danh sách yêu cầu từ người thuê.")
    add_screenshot(doc, "16-admin-dich-vu.png", "Hình 3.12 – Quản lý dịch vụ và yêu cầu bảo trì")
    doc.add_page_break()

    h2(doc, "3.7 Hệ Thống Thông Báo")
    h3(doc, "3.7.1 Danh Sách Thông Báo")
    add_screenshot(doc, "17-admin-thong-bao.png", "Hình 3.13 – Danh sách thông báo (Chung / Riêng)")

    h3(doc, "3.7.2 Tạo Thông Báo")
    body(doc, "Nhập tiêu đề, nội dung. Để trống 'Hợp đồng' = gửi broadcast toàn bộ; "
              "chọn hợp đồng cụ thể = thông báo riêng.")
    add_screenshot(doc, "18-admin-tao-thong-bao.png", "Hình 3.14 – Form tạo thông báo mới")
    doc.add_page_break()

    h2(doc, "3.8 Quản Lý Người Dùng")
    add_screenshot(doc, "19-admin-nguoi-dung.png", "Hình 3.15 – Danh sách người dùng hệ thống")

    h2(doc, "3.9 Quản Lý Khách Hàng")
    add_screenshot(doc, "20-admin-khach-hang.png", "Hình 3.16 – Quản lý khách hàng / người thuê")

    h2(doc, "3.10 Tiện Nghi và Cơ Sở Vật Chất")
    add_screenshot(doc, "21-admin-tien-nghi.png", "Hình 3.17 – Quản lý tiện nghi (WiFi, bãi xe...)")
    add_screenshot(doc, "22-admin-co-so-vat-chat.png", "Hình 3.18 – Quản lý cơ sở vật chất (thang máy, camera...)")
    doc.add_page_break()

    h2(doc, "3.11 Quản Lý Đặt Cọc")
    body(doc, "Ghi nhận giao dịch đặt cọc và hoàn cọc theo từng hợp đồng.")
    add_screenshot(doc, "23-admin-dat-coc.png", "Hình 3.19 – Quản lý giao dịch đặt cọc")

    h2(doc, "3.12 Bảng Giá")
    body(doc, "Quản lý danh mục giá phòng và lịch sử thay đổi giá.")
    add_screenshot(doc, "24-admin-bang-gia.png", "Hình 3.20 – Quản lý bảng giá phòng")

    h2(doc, "3.13 Nhật Ký Hoạt Động")
    body(doc, "Ghi lại toàn bộ hành động của admin: ai làm gì, lúc nào, trên đối tượng nào.")
    add_screenshot(doc, "25-admin-nhat-ky.png", "Hình 3.21 – Nhật ký hoạt động hệ thống")
    doc.add_page_break()

    # ── 4. NGƯỜI THUÊ ────────────────────────────────────────────────────────
    h1(doc, "4. Hướng Dẫn Người Thuê (Customer)")
    body(doc, "Sau khi đăng nhập bằng tài khoản người thuê, hệ thống chuyển đến Dashboard cá nhân.")

    h2(doc, "4.1 Dashboard Người Thuê")
    add_screenshot(doc, "26-customer-dashboard.png", "Hình 4.1 – Dashboard người thuê: tóm tắt hợp đồng và hóa đơn")

    h2(doc, "4.2 Xem Hợp Đồng")
    body(doc, "Xem danh sách hợp đồng với trạng thái Đang hiệu lực / Hết hạn / Đã chấm dứt "
              "và thông tin người cùng thuê.")
    add_screenshot(doc, "27-customer-hop-dong.png", "Hình 4.2 – Danh sách hợp đồng của người thuê")

    h2(doc, "4.3 Xem Hóa Đơn")
    body(doc, "Tra cứu hóa đơn theo tháng với trạng thái Chưa thanh toán / Đã thanh toán / Quá hạn.")
    add_screenshot(doc, "28-customer-hoa-don.png", "Hình 4.3 – Danh sách hóa đơn của người thuê")
    doc.add_page_break()

    h2(doc, "4.4 Yêu Cầu Dịch Vụ")
    body(doc, "Xem danh mục dịch vụ và gửi yêu cầu bảo trì/sửa chữa đến quản trị viên.")
    add_screenshot(doc, "29-customer-dich-vu.png", "Hình 4.4 – Trang dịch vụ của người thuê")

    h2(doc, "4.5 Thông Báo")
    body(doc, "Xem thông báo chung và thông báo riêng dành cho hợp đồng của mình.")
    add_screenshot(doc, "30-customer-thong-bao.png", "Hình 4.5 – Trang thông báo của người thuê")

    h2(doc, "4.6 Hồ Sơ Cá Nhân")
    body(doc, "Cập nhật họ tên, email, số điện thoại, ảnh đại diện và đổi mật khẩu.")
    add_screenshot(doc, "31-customer-ho-so.png", "Hình 4.6 – Trang hồ sơ cá nhân người thuê")
    doc.add_page_break()

    # ── 5. KIẾN TRÚC HỆ THỐNG ──────────────────────────────────────────────
    h1(doc, "5. Kiến Trúc và Công Nghệ")
    info_table(doc, [
        ("Ngôn ngữ backend",  "Java 11"),
        ("Framework",         "Jakarta EE 9.1 (Servlet / JSP)"),
        ("Kiến trúc",         "MVC + DAO Pattern"),
        ("Cơ sở dữ liệu",     "Microsoft SQL Server"),
        ("Frontend",          "Bootstrap 5, HTML/JSP, JavaScript"),
        ("Build tool",        "Apache Maven 3.3+"),
        ("Application server","GlassFish / Apache Tomcat 10+"),
        ("Xác thực",          "Session-based, phân quyền bằng LoginFilter"),
    ])

    h2(doc, "5.1 Phân Quyền Người Dùng")
    info_table(doc, [
        ("Khách (Guest)",       "Xem trang chủ, danh sách phòng, đăng ký/đăng nhập"),
        ("Người thuê (Customer)","Hợp đồng, hóa đơn, dịch vụ, thông báo, hồ sơ cá nhân"),
        ("Quản trị viên (Admin)","Toàn bộ: CRUD phòng, hợp đồng, hóa đơn, người dùng, nhật ký"),
    ])

    # ── Footer note ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    note = doc.add_paragraph(
        "Tài liệu này được soạn thảo phục vụ mục đích hướng dẫn sử dụng và báo cáo học tập. "
        "Ảnh màn hình được ghi lại tự động bằng Playwright tại thời điểm 25/03/2026."
    )
    note.runs[0].font.size = Pt(9)
    note.runs[0].font.italic = True
    note.runs[0].font.color.rgb = DGRAY
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT)
    print(f"Done -> {OUTPUT}")
    print(f"Size: {OUTPUT.stat().st_size / 1024:.0f} KB")

if __name__ == "__main__":
    build()
